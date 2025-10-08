from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from models import Quote, Template
import json
import os
from datetime import datetime

class DocumentGenerator:
    def __init__(self):
        self.output_dir = 'generated_documents'
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_quote_document(self, quote):
        """Generate Word document for quote"""
        try:
            # Create new document
            doc = Document()
            
            # Add header
            self._add_header(doc, quote)
            
            # Add client information
            self._add_client_info(doc, quote)
            
            # Add project details
            self._add_project_details(doc, quote)
            
            # Add AI analysis if available
            if quote.ai_analysis:
                self._add_ai_analysis(doc, quote)
            
            # Add pricing breakdown
            if quote.quote_data:
                self._add_pricing_breakdown(doc, quote)
            
            # Add recommended products
            if quote.recommended_products:
                self._add_recommended_products(doc, quote)
            
            # Add alternative solutions
            if quote.alternative_solutions:
                self._add_alternative_solutions(doc, quote)
            
            # Add terms and conditions
            self._add_terms_conditions(doc)
            
            # Add footer
            self._add_footer(doc, quote)
            
            # Save document
            filename = f"{quote.quote_number.replace(' ', '_')}.docx"
            file_path = os.path.join(self.output_dir, filename)
            doc.save(file_path)
            
            return file_path
            
        except Exception as e:
            print(f"Error generating document: {e}")
            return None
    
    def _add_header(self, doc, quote):
        """Add document header"""
        # Company header
        header_para = doc.add_heading('CCS Structured Cabling Solutions', 0)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Quote title
        title_para = doc.add_heading(f'Quote: {quote.quote_number}', 1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Quote date
        date_para = doc.add_paragraph(f'Date: {quote.created_at.strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Add spacing
    
    def _add_client_info(self, doc, quote):
        """Add client information section"""
        doc.add_heading('Client Information', 2)
        
        client_table = doc.add_table(rows=4, cols=2)
        client_table.style = 'Table Grid'
        client_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Client details
        client_data = [
            ('Client Name:', quote.client_name),
            ('Email:', quote.client_email or 'Not provided'),
            ('Phone:', quote.client_phone or 'Not provided'),
            ('Site Address:', quote.site_address)
        ]
        
        for i, (label, value) in enumerate(client_data):
            client_table.cell(i, 0).text = label
            client_table.cell(i, 1).text = value
        
        doc.add_paragraph()  # Add spacing
    
    def _add_project_details(self, doc, quote):
        """Add project details section"""
        doc.add_heading('Project Details', 2)
        
        # Project title
        doc.add_paragraph(f'Project: {quote.project_title}')
        
        # Project description
        if quote.project_description:
            doc.add_paragraph(f'Description: {quote.project_description}')
        
        # Building information
        building_info = []
        if quote.building_type:
            building_info.append(f'Building Type: {quote.building_type}')
        if quote.building_size:
            building_info.append(f'Building Size: {quote.building_size} sqm')
        if quote.number_of_floors:
            building_info.append(f'Number of Floors: {quote.number_of_floors}')
        if quote.number_of_rooms:
            building_info.append(f'Number of Rooms: {quote.number_of_rooms}')
        
        for info in building_info:
            doc.add_paragraph(info)
        
        # Requirements
        doc.add_paragraph('Requirements:')
        requirements = []
        if quote.wifi_requirements:
            requirements.append('• WiFi Installation')
        if quote.cctv_requirements:
            requirements.append('• CCTV System')
        if quote.door_entry_requirements:
            requirements.append('• Door Entry System')
        if quote.cabling_type:
            requirements.append(f'• {quote.cabling_type.upper()} Cabling')
        
        for req in requirements:
            doc.add_paragraph(req, style='List Bullet')
        
        # Special requirements
        if quote.special_requirements:
            doc.add_paragraph(f'Special Requirements: {quote.special_requirements}')
        
        doc.add_paragraph()  # Add spacing
    
    def _add_ai_analysis(self, doc, quote):
        """Add AI analysis section"""
        doc.add_heading('Technical Analysis', 2)
        doc.add_paragraph(quote.ai_analysis)
        doc.add_paragraph()  # Add spacing
    
    def _add_pricing_breakdown(self, doc, quote):
        """Add pricing breakdown section"""
        doc.add_heading('Pricing Breakdown', 2)
        
        try:
            quote_data = json.loads(quote.quote_data)
            
            # Materials section
            if quote_data.get('materials'):
                doc.add_heading('Materials', 3)
                
                materials_table = doc.add_table(rows=1, cols=5)
                materials_table.style = 'Table Grid'
                materials_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header row
                hdr_cells = materials_table.rows[0].cells
                hdr_cells[0].text = 'Item'
                hdr_cells[1].text = 'Part Number'
                hdr_cells[2].text = 'Quantity'
                hdr_cells[3].text = 'Unit Price'
                hdr_cells[4].text = 'Total'
                
                # Add material rows
                has_estimated = False
                for material in quote_data['materials']:
                    row_cells = materials_table.add_row().cells
                    
                    # Add asterisk for estimated pricing
                    item_name = material.get('item', '')
                    unit_price = material.get('unit_price', 0)
                    total_price = material.get('total', 0)
                    is_estimated = material.get('is_estimated', False)
                    
                    if is_estimated:
                        has_estimated = True
                        item_name += " *"
                    
                    row_cells[0].text = item_name
                    row_cells[1].text = material.get('part_number', '')
                    row_cells[2].text = f"{material.get('quantity', 0)} {material.get('unit', '')}"
                    row_cells[3].text = f"£{unit_price:.2f}{' *' if is_estimated else ''}"
                    row_cells[4].text = f"£{total_price:.2f}{' *' if is_estimated else ''}"
                
                doc.add_paragraph(f"Total Materials: £{quote_data.get('total_materials', 0):.2f}")
                
                # Add pricing legend if there are estimated prices
                if has_estimated:
                    legend_para = doc.add_paragraph()
                    legend_para.add_run("Pricing Legend: ").bold = True
                    legend_para.add_run("* = Estimated pricing based on average market data. Actual supplier quotes may vary. Please confirm final pricing before proceeding.")
                    legend_para.runs[0].font.size = Pt(9)
                    legend_para.runs[1].font.size = Pt(9)
            
            # Labor section
            if quote_data.get('labor'):
                doc.add_heading('Labor', 3)
                
                labor_table = doc.add_table(rows=1, cols=4)
                labor_table.style = 'Table Grid'
                labor_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header row
                hdr_cells = labor_table.rows[0].cells
                hdr_cells[0].text = 'Service'
                hdr_cells[1].text = 'Hours'
                hdr_cells[2].text = 'Rate'
                hdr_cells[3].text = 'Total'
                
                # Add labor rows
                for labor in quote_data['labor']:
                    row_cells = labor_table.add_row().cells
                    row_cells[0].text = labor.get('item', '')
                    row_cells[1].text = f"{labor.get('quantity', 0)} {labor.get('unit', '')}"
                    row_cells[2].text = f"£{labor.get('unit_price', 0):.2f}"
                    row_cells[3].text = f"£{labor.get('total', 0):.2f}"
                
                doc.add_paragraph(f"Total Labor: £{quote_data.get('total_labor', 0):.2f}")
            
            # Total cost
            doc.add_paragraph()
            total_para = doc.add_paragraph(f"TOTAL COST: £{quote_data.get('total_cost', 0):.2f}")
            total_para.runs[0].bold = True
            total_para.runs[0].font.size = Pt(14)
            
        except Exception as e:
            doc.add_paragraph(f"Error displaying pricing: {str(e)}")
        
        doc.add_paragraph()  # Add spacing
    
    def _add_recommended_products(self, doc, quote):
        """Add recommended products section"""
        try:
            products = json.loads(quote.recommended_products)
            
            if products:
                doc.add_heading('Recommended Products', 2)
                
                for product in products:
                    doc.add_paragraph(f"• {product}")
                
                doc.add_paragraph()  # Add spacing
                
        except Exception as e:
            print(f"Error adding recommended products: {e}")
    
    def _add_alternative_solutions(self, doc, quote):
        """Add alternative solutions section"""
        try:
            alternatives = json.loads(quote.alternative_solutions)
            
            if alternatives:
                doc.add_heading('Alternative Solutions', 2)
                
                for i, alternative in enumerate(alternatives, 1):
                    doc.add_heading(f'Option {i}', 3)
                    doc.add_paragraph(str(alternative))
                
                doc.add_paragraph()  # Add spacing
                
        except Exception as e:
            print(f"Error adding alternative solutions: {e}")
    
    def _add_terms_conditions(self, doc):
        """Add terms and conditions"""
        doc.add_heading('Terms and Conditions', 2)
        
        terms = [
            "• This quote is valid for 30 days from the date of issue.",
            "• Payment terms: 50% deposit on acceptance, 50% on completion.",
            "• Installation timeline will be confirmed upon order acceptance.",
            "• All materials are covered by manufacturer warranty.",
            "• Installation includes 12 months warranty on workmanship.",
            "• Any changes to the scope of work may affect pricing and timeline.",
            "• Client is responsible for providing safe access to all work areas.",
            "• Quote excludes any structural modifications or building works."
        ]
        
        for term in terms:
            doc.add_paragraph(term, style='List Bullet')
        
        doc.add_paragraph()  # Add spacing
    
    def _add_footer(self, doc, quote):
        """Add document footer"""
        doc.add_paragraph()
        doc.add_paragraph("Thank you for considering CCS Structured Cabling Solutions.")
        doc.add_paragraph()
        
        # Contact information
        contact_info = [
            "CCS Structured Cabling Solutions",
            "Email: info@ccs-cabling.com",
            "Phone: [Your Phone Number]",
            "Website: [Your Website]"
        ]
        
        for info in contact_info:
            doc.add_paragraph(info)
    
    def generate_email_template(self, quote):
        """Generate email template for quote"""
        try:
            template = f"""
Subject: Quote {quote.quote_number} - {quote.project_title}

Dear {quote.client_name},

Thank you for your interest in our structured cabling services. Please find attached your detailed quote for the {quote.project_title} project.

Quote Details:
- Quote Number: {quote.quote_number}
- Project: {quote.project_title}
- Site: {quote.site_address}
- Estimated Cost: £{quote.estimated_cost:.2f} (if available)

The attached document contains:
- Detailed project analysis
- Recommended equipment and materials
- Complete pricing breakdown
- Alternative solutions (if applicable)
- Terms and conditions

This quote is valid for 30 days. If you have any questions or would like to discuss any aspect of the quote, please don't hesitate to contact us.

We look forward to working with you on this project.

Best regards,
CCS Structured Cabling Solutions
"""
            
            return template
            
        except Exception as e:
            print(f"Error generating email template: {e}")
            return None

